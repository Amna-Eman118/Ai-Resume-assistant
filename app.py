import io
import json
import os
from typing import List

import streamlit as st
from google import genai
from google.genai import types
from pydantic import BaseModel, Field
from pypdf import PdfReader
from docx import Document


MODEL_NAME = "gemini-1.5-flash"


class Improvement(BaseModel):
    priority: str = Field(description="High, Medium, or Low")
    issue: str = Field(description="Specific resume issue")
    recommendation: str = Field(description="Concrete improvement")
    example: str = Field(description="A concise example of improved wording")


class ResumeAnalysis(BaseModel):
    ats_score: int = Field(description="Overall ATS-readiness score from 0 to 100")
    summary: str = Field(description="Short explanation of the score")
    strengths: List[str] = Field(description="3 to 6 strongest ATS-friendly aspects")
    improvements: List[Improvement] = Field(description="5 to 10 concrete improvements")
    missing_keywords: List[str] = Field(description="Important keywords missing or weakly represented")
    section_feedback: List[str] = Field(description="Feedback on resume sections and formatting")
    keyword_match_notes: List[str] = Field(description="Notes about alignment with the supplied job description")


def get_api_key() -> str:
    """Read the Gemini key from Streamlit secrets first, then environment variables."""
    try:
        key = st.secrets.get("GEMINI_API_KEY", "")
    except Exception:
        key = ""
    return key or os.getenv("GEMINI_API_KEY", "")


def extract_resume_text(uploaded_file) -> str:
    """Extract text from PDF, DOCX, or TXT uploads."""
    file_bytes = uploaded_file.getvalue()
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(file_bytes))
        pages = []
        for page in reader.pages:
            pages.append(page.extract_text() or "")
        return "\n\n".join(pages).strip()

    if filename.endswith(".docx"):
        document = Document(io.BytesIO(file_bytes))
        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                parts.append(" | ".join(cell.text.strip() for cell in row.cells))
        return "\n".join(parts).strip()

    if filename.endswith(".txt"):
        return file_bytes.decode("utf-8", errors="replace").strip()

    raise ValueError("Unsupported file type. Please upload PDF, DOCX, or TXT.")


def analyze_resume(resume_text: str, job_description: str) -> ResumeAnalysis:
    api_key = get_api_key()
    if not api_key:
        raise RuntimeError(
            "Gemini API key not found. Add GEMINI_API_KEY to Streamlit Secrets "
            "or set it as an environment variable."
        )

    client = genai.Client(api_key=api_key)

    jd_text = job_description.strip() or "No job description supplied. Evaluate general ATS readiness."

    prompt = f"""
You are an expert ATS resume reviewer and recruiter.

Analyze the resume below for ATS-readiness. Give an ATS score from 0 to 100.
This is an estimated readiness score, NOT a guarantee of how any specific ATS will score it.

Evaluate:
1. Standard section headings and organization
2. ATS-readable formatting and likely parsing problems
3. Keyword relevance
4. Skills section quality
5. Work experience quality and measurable achievements
6. Education/certifications
7. Contact information completeness
8. Action verbs and concise wording
9. Keyword alignment with the job description when supplied
10. Common ATS risks such as tables, columns, graphics, icons, headers/footers, unusual symbols, or overly decorative formatting when they are evident from the extracted text

Important:
- Do not invent experience, skills, education, metrics, employers, or certifications.
- Recommendations must be truthful and based on information actually present or clearly missing.
- If a job description is supplied, identify relevant missing/weak keywords without keyword stuffing.
- Explain that visual formatting cannot always be fully detected from extracted text.
- Keep examples concise and clearly label them as examples/templates, not facts.

JOB DESCRIPTION:
{jd_text}

RESUME:
{resume_text}
"""

    response = client.models.generate_content(
        model=MODEL_NAME,
        contents=prompt,
        config=types.GenerateContentConfig(
            temperature=0.2,
            response_mime_type="application/json",
            response_schema=ResumeAnalysis,
        ),
    )

    if not response.text:
        raise RuntimeError("Gemini returned an empty response.")

    return ResumeAnalysis.model_validate_json(response.text)


def render_analysis(result: ResumeAnalysis):
    st.subheader("ATS Score")
    st.progress(result.ats_score / 100)
    st.metric("Estimated ATS Score", f"{result.ats_score}/100")

    st.write(result.summary)

    col1, col2 = st.columns(2)
    with col1:
        st.subheader("Strengths")
        for item in result.strengths:
            st.success(item)

    with col2:
        st.subheader("Missing / Weak Keywords")
        if result.missing_keywords:
            for item in result.missing_keywords:
                st.warning(item)
        else:
            st.info("No major missing keywords were identified.")

    st.subheader("Improvements")
    for i, item in enumerate(result.improvements, start=1):
        with st.expander(f"{i}. {item.issue} — {item.priority} priority"):
            st.write(f"**Recommendation:** {item.recommendation}")
            st.write(f"**Example:** {item.example}")

    st.subheader("Section & Formatting Feedback")
    for item in result.section_feedback:
        st.write(f"- {item}")

    if result.keyword_match_notes:
        st.subheader("Job Description Match Notes")
        for item in result.keyword_match_notes:
            st.write(f"- {item}")


st.set_page_config(
    page_title="Resume ATS Analyzer",
    page_icon="📄",
    layout="wide",
)

st.title("📄 Resume ATS Analyzer")
st.caption(
    "Upload your resume to get an estimated ATS-readiness score, keyword feedback, "
    "and practical improvement suggestions powered by Gemini Flash."
)

with st.sidebar:
    st.header("How it works")
    st.write("1. Upload PDF, DOCX, or TXT")
    st.write("2. Optionally paste a job description")
    st.write("3. Gemini analyzes ATS readiness")
    st.write("4. Review the score and improvements")
    st.info(
        "Tip: For the best keyword feedback, paste the exact job description you are applying to."
    )

uploaded_file = st.file_uploader(
    "Upload your resume",
    type=["pdf", "docx", "txt"],
    help="PDF, DOCX, and TXT are supported.",
)

job_description = st.text_area(
    "Job Description (optional)",
    height=220,
    placeholder="Paste the job description here for keyword and role alignment analysis...",
)

if uploaded_file:
    st.caption(f"Selected: {uploaded_file.name}")

    if st.button("🔍 Analyze Resume", type="primary", use_container_width=True):
        try:
            with st.spinner("Extracting resume text and analyzing with Gemini..."):
                resume_text = extract_resume_text(uploaded_file)

                if not resume_text:
                    st.error(
                        "No readable text was found. If this is a scanned/image-only PDF, "
                        "please use a text-based PDF or DOCX."
                    )
                    st.stop()

                # Keep prompts reasonably sized while preserving the beginning and end.
                max_chars = 50000
                if len(resume_text) > max_chars:
                    resume_text = (
                        resume_text[:35000]
                        + "\n\n[Middle of resume truncated for processing]\n\n"
                        + resume_text[-15000:]
                    )

                result = analyze_resume(resume_text, job_description)
                st.session_state["analysis"] = result
                st.session_state["resume_name"] = uploaded_file.name

        except Exception as exc:
            st.error(f"Analysis failed: {exc}")

if "analysis" in st.session_state:
    st.divider()
    st.caption(f"Analysis for: {st.session_state.get('resume_name', 'Resume')}")
    render_analysis(st.session_state["analysis"])
else:
    st.info("Upload a resume and click **Analyze Resume** to get started.")
